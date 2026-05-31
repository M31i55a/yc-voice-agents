#!/usr/bin/env python3
"""
Generate a comprehensive tutorial document for the Bayview Pharmacy Voice Agent.
Enhanced version with project structure, frontend/backend explanation, and Docker.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ============================================================================
# TITLE AND INTRODUCTION
# ============================================================================

title = doc.add_paragraph()
title_run = title.add_run('Bayview Pharmacy Voice Agent')
title_run.bold = True
title_run.font.size = Pt(24)
title_run.font.color.rgb = RGBColor(0, 102, 204)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('A Comprehensive Tutorial for Beginners')
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(14)
subtitle_run.italic = True
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

date_para = doc.add_paragraph('May 2026 • Complete Edition with Architecture & Deployment')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # Spacing

# ============================================================================
# TABLE OF CONTENTS
# ============================================================================

doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Introduction and Overview',
    '2. The Problem & Solution',
    '3. Project Structure & File Tree',
    '4. Frontend Architecture (demo_client)',
    '5. Backend Architecture (server)',
    '6. How Frontend and Backend Communicate',
    '7. Pipeline & Services Deep Dive',
    '8. The Tools System & Verification',
    '9. Self-Healing Loop (Cekura)',
    '10. Running Locally',
    '11. Docker & Containerization',
    '12. Deploying to Production',
    '13. Extending the Project',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============================================================================
# SECTION 1: INTRODUCTION
# ============================================================================

doc.add_heading('1. Introduction and Overview', level=1)

doc.add_heading('What is Bayview Pharmacy?', level=2)
doc.add_paragraph(
    'Bayview Pharmacy is an AI voice agent that helps patients refill prescriptions and check medication status. '
    'It combines speech recognition, AI language understanding, natural speech generation, and optional video '
    'with an AI avatar to create an accessible, human-like pharmacy experience available 24/7.'
)

doc.add_paragraph('Instead of calling a busy pharmacy and waiting on hold, patients can:')
doc.add_paragraph('Call a phone number or open a web app', style='List Bullet')
doc.add_paragraph('Speak naturally in English or Spanish', style='List Bullet')
doc.add_paragraph('See an AI avatar if using video', style='List Bullet')
doc.add_paragraph('Verify identity and manage prescriptions in minutes', style='List Bullet')

doc.add_heading('Key Capabilities', level=2)
doc.add_paragraph('🎥 Video with AI Avatar: Google Meet-style UI with gesture recognition', style='List Bullet')
doc.add_paragraph('🗣️ Multilingual: Automatic English/Spanish detection and switching', style='List Bullet')
doc.add_paragraph('🔐 Secure: Identity verification ENFORCED in code before any data access', style='List Bullet')
doc.add_paragraph('📱 Multi-transport: Phone (Twilio), WebRTC (web/local), Pipecat Cloud', style='List Bullet')
doc.add_paragraph('🤖 Self-improving: Cekura evaluation loop finds and fixes failures automatically', style='List Bullet')

doc.add_page_break()

# ============================================================================
# SECTION 2: THE PROBLEM & SOLUTION
# ============================================================================

doc.add_heading('2. The Problem & Solution', level=1)

doc.add_heading('The Real Problem', level=2)
doc.add_paragraph(
    'Prescription refill calls are a bottleneck. Patients struggle with:'
)
doc.add_paragraph('Long hold times (5-10+ minutes)', style='List Bullet')
doc.add_paragraph('Repeating sensitive info multiple times', style='List Bullet')
doc.add_paragraph('Accessibility barriers (hearing loss, speech disabilities, non-English speakers)', style='List Bullet')
doc.add_paragraph('No 24/7 availability', style='List Bullet')
doc.add_paragraph('Difficulty understanding medication instructions over audio alone', style='List Bullet')

doc.add_heading('Bayview\'s Solution', level=2)
doc.add_paragraph(
    'An AI agent that:'
)
doc.add_paragraph('Answers immediately (no hold time)', style='List Bullet')
doc.add_paragraph('Supports video so patients can see lip movement and facial expressions', style='List Bullet')
doc.add_paragraph('Handles multilingual requests automatically', style='List Bullet')
doc.add_paragraph('Uses gesture recognition to improve accessibility', style='List Bullet')
doc.add_paragraph('Available 24/7, reducing pharmacy staff workload', style='List Bullet')

doc.add_page_break()

# ============================================================================
# SECTION 3: PROJECT STRUCTURE
# ============================================================================

doc.add_heading('3. Project Structure & File Tree', level=1)

doc.add_heading('Complete Project Layout', level=2)

project_structure = """yc-voice-agents/
│
├─ server/                          # Backend: The AI agent & services
│  ├─ bot-nemotron.py               # PRIMARY: Nemotron LLM + NVIDIA STT + Gradium TTS
│  ├─ bot-gpt.py                    # FALLBACK: OpenAI GPT-4 + Gradium STT/TTS
│  ├─ mock_backend.py               # Patient data (in-memory dict, replace with DB)
│  ├─ nemotron_llm.py               # NVIDIA Nemotron LLM service wrapper
│  ├─ nvidia_stt.py                 # Custom NVIDIA Parakeet speech-to-text
│  ├─ stt_provider.py               # STT service factory (chooses Gradium or NVIDIA)
│  ├─ language_router.py            # Detects English/Spanish from speech + gestures
│  ├─ video_avatar.py               # AI avatar rendering for video calls
│  │
│  ├─ demo_client/                  # Frontend: Browser UI & WebRTC client
│  │  ├─ index.html                 # Pre-join + in-call UI (Google Meet style)
│  │  ├─ app.js                     # Main logic: WebRTC, RTVI, MediaPipe
│  │  ├─ styles.css                 # UI styling (light/dark mode, responsive)
│  │
│  ├─ Dockerfile                    # Container config (builds to Pipecat Cloud)
│  ├─ pyproject.toml                # Python dependencies (uv)
│  ├─ pcc-deploy.toml               # Pipecat Cloud deployment config
│
├─ harness/                         # Self-improvement: Cekura → GPT → redeploy
│  ├─ self_heal.py                  # Asks GPT to fix failed scenarios
│  ├─ webhook_server.py             # Listens for Cekura failures
│  ├─ generate_dashboard.py         # Creates KPI dashboard from Cekura data
│  ├─ serve_dashboard.py            # Runs dashboard server with live refresh
│  ├─ examples/
│  │  └─ bayview_cekura_report_sample.json  # Sample Cekura eval results
│  ├─ tests/
│  │  ├─ test_self_heal.py
│  │  └─ test_webhook_server.py
│
├─ docs/                            # Design docs & specifications
│  └─ superpowers/
│     ├─ specs/
│     │  └─ 2026-05-30-self-healing-infra-design.md
│     └─ plans/
│        └─ 2026-05-30-self-healing-infra.md
│
├─ README.md                        # Main project README
├─ CLAUDE.md                        # Hackathon runbook with gotchas & conventions
├─ CEKURA_RUNBOOK.md               # Cekura evaluation guide"""

code_para = doc.add_paragraph(project_structure)
for run in code_para.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_heading('What Each Directory Does', level=2)

directories = [
    ('server/', 
     'The backend agent and web UI. Everything runs from here during development. Python + Pipecat.'),
    
    ('server/demo_client/', 
     'Browser UI (HTML/CSS/JS). This is what callers see on their screen—camera preview, call controls, transcript.'),
    
    ('harness/', 
     'Automated evaluation and fixing. Listens to Cekura, proposes fixes via GPT, redeploys if tests improve.'),
    
    ('docs/', 
     'Design documents. Explains the architecture decisions and self-healing infrastructure.'),
]

for dir_name, description in directories:
    doc.add_heading(dir_name, level=3)
    doc.add_paragraph(description)

doc.add_page_break()

# ============================================================================
# SECTION 4: FRONTEND ARCHITECTURE
# ============================================================================

doc.add_heading('4. Frontend Architecture (demo_client)', level=1)

doc.add_heading('What is the Frontend?', level=2)
doc.add_paragraph(
    'The frontend is a single-page web application (SPA) that runs in the browser. It\'s a Google Meet-style UI '
    'where patients can join a call with the pharmacy bot, see a camera preview, adjust microphone/camera, and interact in real-time.'
)

doc.add_heading('Frontend Tech Stack', level=2)
doc.add_paragraph('HTML / CSS / JavaScript (no React/Vue—vanilla for simplicity)', style='List Bullet')
doc.add_paragraph('WebRTC (peer-to-peer audio/video between browser and backend)', style='List Bullet')
doc.add_paragraph('RTVI (Real Time Voice Interaction protocol—custom Pipecat protocol)', style='List Bullet')
doc.add_paragraph('MediaPipe Tasks Vision (gesture + object detection for accessibility)', style='List Bullet')

doc.add_heading('Key Files', level=2)

frontend_files = [
    ('index.html',
     'Defines the UI structure: pre-join screen (camera check, device selection), in-call screen (video tiles, transcript, controls).'),
    
    ('app.js (~1500 lines)',
     'The main logic: WebRTC connection setup, RTVI protocol handling, audio level metering, MediaPipe vision integration.'),
    
    ('styles.css',
     'Responsive styling for mobile/tablet/desktop. Light/dark theme. Google Material Design color scheme.'),
]

for filename, description in frontend_files:
    doc.add_heading(filename, level=3)
    doc.add_paragraph(description)

doc.add_heading('User Journey (Frontend Flow)', level=2)

journey_steps = [
    ('User opens pharmacy.example.com',
     'Browser loads index.html + app.js + styles.css'),
    
    ('Pre-join screen appears',
     'User sees camera preview, selects microphone/camera/speaker, can test audio levels (level meter shown)'),
    
    ('User clicks "Connect"',
     'app.js initiates WebRTC negotiation with backend server (SmallWebRTC transport)'),
    
    ('Call connects',
     'MediaPipe starts analyzing camera for gestures + objects (cup/bottle for "show me") in background'),
    
    ('In-call UI appears',
     'User sees their video, agent video (if video avatar enabled), live transcript, mute/camera buttons'),
    
    ('Agent speaks',
     'Audio frames come from backend, displayed as live captions (STT transcript), user can hear TTS audio'),
    
    ('User speaks',
     'Browser captures audio, sends to backend via WebRTC data channel + audio track'),
    
    ('Gesture detected',
     'e.g., user points up (English) or shows victory sign (Spanish). Frontend sends language hint to backend.'),
    
    ('Call ends',
     'User clicks hang up, WebRTC connection closes, transcript saved'),
]

for i, (step, detail) in enumerate(journey_steps, 1):
    p = doc.add_paragraph(f'{i}. {step}')
    p.style = 'List Number'
    doc.add_paragraph(detail, style='List Bullet 2')

doc.add_heading('MediaPipe Integration (Accessibility)', level=2)
doc.add_paragraph(
    'The frontend runs MediaPipe object detection + gesture recognition in the browser, entirely on the user\'s device '
    '(no server-side vision processing):'
)

doc.add_paragraph(
    'Object Detection: Detects cups/bottles. If user shows a cup-like object, a visual cue appears ("I see a cup—tap to request refill help")',
    style='List Bullet')

doc.add_paragraph(
    'Gesture Recognition: Detects hand poses (Victory sign, Pointing Up). Sends language hints to backend '
    '(e.g., "user showed victory sign → Spanish").',
    style='List Bullet')

doc.add_paragraph(
    'Privacy: All analysis happens locally. No video is sent to backend. Only text hints are sent.',
    style='List Bullet')

doc.add_page_break()

# ============================================================================
# SECTION 5: BACKEND ARCHITECTURE
# ============================================================================

doc.add_heading('5. Backend Architecture (server)', level=1)

doc.add_heading('What is the Backend?', level=2)
doc.add_paragraph(
    'The backend is a Python application built with Pipecat that orchestrates the voice call. It:'
)
doc.add_paragraph('Receives audio from the frontend via WebRTC', style='List Bullet')
doc.add_paragraph('Converts speech to text (STT)', style='List Bullet')
doc.add_paragraph('Routes to language processor (English or Spanish)', style='List Bullet')
doc.add_paragraph('Feeds text + context to the LLM', style='List Bullet')
doc.add_paragraph('LLM calls tools (verify_identity, get_prescriptions, refill_prescription)', style='List Bullet')
doc.add_paragraph('Tools query mock_backend.py for patient data', style='List Bullet')
doc.add_paragraph('LLM generates response, converted to speech (TTS)', style='List Bullet')
doc.add_paragraph('Audio sent back to frontend', style='List Bullet')

doc.add_heading('Backend Tech Stack', level=2)
doc.add_paragraph('Python 3.11+ (async/await throughout)', style='List Bullet')
doc.add_paragraph('Pipecat (voice orchestration framework)', style='List Bullet')
doc.add_paragraph('Nemotron 3 Super 120B (primary LLM) or GPT-4 (fallback)', style='List Bullet')
doc.add_paragraph('NVIDIA Parakeet (optional STT) or Gradium (default STT)', style='List Bullet')
doc.add_paragraph('Gradium TTS (text-to-speech)', style='List Bullet')
doc.add_paragraph('FastAPI (for WebSocket endpoints)', style='List Bullet')

doc.add_heading('Entry Points', level=2)

entry_points = [
    ('bot-nemotron.py (PRIMARY)',
     'Uses NVIDIA Nemotron LLM + optional NVIDIA Parakeet STT. Requires NVIDIA endpoints. Fastest, open-weight.'),
    
    ('bot-gpt.py (FALLBACK)',
     'Uses OpenAI GPT-4.1 LLM + Gradium STT. Good fallback if NVIDIA is unavailable. Requires OpenAI key.'),
]

for name, desc in entry_points:
    doc.add_heading(name, level=3)
    doc.add_paragraph(desc)

doc.add_heading('Core Services (Inside bot-nemotron.py)', level=2)

services_desc = [
    ('STT Service',
     'Nemotron Speech (WebSocket ASR) or Gradium. Converts caller audio → text in real time.'),
    
    ('Language Router',
     'Custom processor (language_router.py). Detects English/Spanish from text + gesture hints. Injects language context into LLM.'),
    
    ('LLM Service',
     'Nemotron 3 Super 120B via vLLM (nemotron_llm.py). Receives text + context, outputs text + tool calls.'),
    
    ('Tool Functions',
     'Python async functions: verify_identity(), get_prescriptions(), refill_prescription(), end_call(). '
     'LLM calls these, they return results that feed back to LLM.'),
    
    ('TTS Service',
     'Gradium TTS. Converts LLM response text → audio stream. Multi-voice support.'),
    
    ('Avatar Service (optional)',
     'video_avatar.py. Renders AI face with lip-sync. Only on video calls; disabled for Twilio phone.'),
]

for service, description in services_desc:
    doc.add_heading(service, level=3)
    doc.add_paragraph(description)

doc.add_heading('Data Flow Inside the Backend', level=2)

data_flow = """1. Caller speaks: "I need to refill my Lisinopril"
   ↓
2. STT (NVIDIA Parakeet/Gradium): "I need to refill my Lisinopril"
   ↓
3. Language Router: Detects English, adds {language: "English"} to context
   ↓
4. LLM receives: Text + Context + Available Tools
   LLM thinks: "User wants refill, but I don't know if they're verified yet"
   LLM calls: verify_identity(full_name="...", date_of_birth="...")
   ↓
5. Tool Handler: verify_identity() checks mock_backend.py
   Returns: {verified: true} or {verified: false, failed_attempts: 1}
   ↓
6. LLM receives tool result, continues conversation
   If verified: "Great! I can refill Lisinopril. Confirm pickup in 2 days?"
   ↓
7. LLM output: "Great! I can refill..."
   ↓
8. TTS (Gradium): Generates audio of that sentence
   ↓
9. Audio sent back to caller via WebRTC"""

flow_para = doc.add_paragraph(data_flow)
for run in flow_para.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_page_break()

# ============================================================================
# SECTION 6: FRONTEND-BACKEND COMMUNICATION
# ============================================================================

doc.add_heading('6. How Frontend and Backend Communicate', level=1)

doc.add_heading('Communication Protocols', level=2)

doc.add_heading('1. WebRTC (Audio/Video)', level=3)
doc.add_paragraph(
    'Real-time media transport. Frontend browser ↔ Backend Pipecat. '
    'Handles audio & video frames with low latency. Uses DTLS encryption.'
)

doc.add_heading('2. RTVI Protocol (Data Channel)', level=3)
doc.add_paragraph(
    'Custom protocol over WebRTC data channel. Frontend sends:'
)
doc.add_paragraph('Gesture detection events (e.g., {"gesture": "victory", "language": "es"})', style='List Bullet')
doc.add_paragraph('User status (muted, camera on/off)', style='List Bullet')
doc.add_paragraph('Transcript requests', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Backend sends back:')
doc.add_paragraph('Transcript updates (agent speech)', style='List Bullet')
doc.add_paragraph('LLM context (available tools)', style='List Bullet')
doc.add_paragraph('Bot status (speaking, waiting)', style='List Bullet')

doc.add_heading('The Three-Stage Connection', level=2)

connection_stages = [
    ('Pre-join (Frontend Only)',
     'User in browser, testing camera/mic. No server connection yet. MediaPipe is warming up. '
     'User clicks "Connect"'),
    
    ('SDP Offer/Answer (WebRTC Negotiation)',
     'Browser generates WebRTC offer → Backend receives → generates answer → '
     'Browser gets answer. ICE candidates exchanged (STUN/TURN). '
     'Connection established.'),
    
    ('Connected (Streaming)',
     'WebRTC connection is live. Audio/video flowing. RTVI messages exchanged. '
     'Pipecat pipeline processes frames. Conversation happens in real-time.'),
]

for i, (stage, description) in enumerate(connection_stages, 1):
    doc.add_heading(f'{i}. {stage}', level=3)
    doc.add_paragraph(description)

doc.add_heading('Diagram: Frontend ↔ Backend', level=2)

diagram = """Browser (Frontend)                  Network                  Server (Backend)
┌─────────────────────────┐                            ┌──────────────────────┐
│  index.html + app.js    │                            │   bot-nemotron.py    │
│  ┌───────────────────┐  │                            │  ┌────────────────┐  │
│  │ Pre-join Screen   │  │                            │  │  Pipecat      │  │
│  │ (camera check)    │  │                            │  │  Pipeline     │  │
│  └───────────────────┘  │                            │  │              │  │
│           ↓             │                            │  │  ┌────────┐   │  │
│  ┌───────────────────┐  │  WebRTC Offer   ────────>  │  │  │  STT   │   │  │
│  │  app.js initiate  │  │                            │  │  └────────┘   │  │
│  │  WebRTC.connect() │  │  WebRTC Answer   <──────  │  │  ┌────────┐   │  │
│  └───────────────────┘  │                            │  │  │  LLM   │   │  │
│           ↓             │                            │  │  └────────┘   │  │
│  ┌───────────────────┐  │  Audio Frames   ←────────  │  │  ┌────────┐   │  │
│  │  In-call Screen   │  │                            │  │  │  TTS   │   │  │
│  │ (call, transcript)│  │  Audio Frames   ────────>  │  │  └────────┘   │  │
│  └───────────────────┘  │                            │  │              │  │
│           ↓             │  RTVI Messages  ←─────────  │  │  mock_backend│  │
│  ┌───────────────────┐  │                            │  │  (patient DB)│  │
│  │  MediaPipe Vision │  │  Gesture Hints  ────────>  │  └────────────┘  │
│  │  (gesture detect) │  │                            │                  │
│  └───────────────────┘  │                            │                  │
└─────────────────────────┘                            └──────────────────────┘"""

diagram_para = doc.add_paragraph(diagram)
for run in diagram_para.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(8)

doc.add_heading('Specific Messages Exchanged', level=2)

doc.add_heading('Frontend → Backend (Common RTVI Messages)', level=3)

frontend_msgs = [
    '{role: "user", content: "I need a refill"}  # User speech transcript',
    '{gesture_detected: true, gesture: "victory"}  # User made victory sign',
    '{language_hint: "es"}  # Gesture or menu selection hints Spanish',
    '{muted: true}  # User muted mic',
    '{camera_enabled: false}  # User turned off camera',
]

for msg in frontend_msgs:
    doc.add_paragraph(msg, style='List Bullet')

doc.add_heading('Backend → Frontend (Common RTVI Messages)', level=3)

backend_msgs = [
    '{role: "assistant", content: "Great, I can help with that"}  # Agent response',
    '{role: "assistant", tool_calls: [{name: "get_prescriptions", args: {...}}]}  # Tool invocation',
    '{role: "tool", result: {prescriptions: [...]}}  # Tool result',
    '{bot_speaking: true}  # Agent is currently speaking',
    '{transcript_updated: true}  # New caption available',
]

for msg in backend_msgs:
    doc.add_paragraph(msg, style='List Bullet')

doc.add_page_break()

# ============================================================================
# SECTION 7: PIPELINE & SERVICES
# ============================================================================

doc.add_heading('7. Pipeline & Services Deep Dive', level=1)

doc.add_heading('The Pipecat Pipeline', level=2)

doc.add_paragraph(
    'Pipecat is a framework that chains together processors and services. Think of it as a factory assembly line.'
)

pipeline_visual = """SmallWebRTC Transport
        ↓
   [Audio In]
        ↓
   Silero VAD (Voice Activity Detection)
        ↓
   STT Service (NVIDIA Parakeet or Gradium)
        ↓
   [Text Frames]
        ↓
   Language Router (detects English/Spanish)
        ↓
   LLM Context Aggregator (builds conversation context)
        ↓
   LLM Service (Nemotron 3 Super or GPT-4)
        ↓
   LLM calls Tool Functions
        ↓
   Tool Functions execute (verify_identity, get_prescriptions, etc.)
        ↓
   Tool Results fed back to LLM
        ↓
   LLM generates Response Text
        ↓
   TTS Service (Gradium)
        ↓
   [Audio Out]
        ↓
   Optional: Avatar Service (adds face overlay)
        ↓
   SmallWebRTC Transport (sends to browser)"""

pipeline_para = doc.add_paragraph(pipeline_visual)
for run in pipeline_para.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_heading('Each Pipeline Component', level=2)

components = [
    ('SmallWebRTC Transport',
     'Connects to browser. Exchanges WebRTC offers/answers. Streams audio/video frames to/from pipeline.'),
    
    ('Silero VAD',
     'Voice Activity Detection. Listens for silence. Tells the pipeline when the user stops talking '
     '(so the bot knows when to respond).'),
    
    ('STT Service',
     'Converts audio frames → text. Async, streaming. Gradium can add captions in real-time. '
     'NVIDIA Parakeet is more accurate but requires NVIDIA endpoints.'),
    
    ('Language Router',
     'Custom Pipecat processor. Inspects transcription + gesture hints. '
     'Injects {language: "English"} or {language: "Spanish"} into LLM context.'),
    
    ('LLM Context Aggregator',
     'Builds the "conversation context" (list of prior messages + current state). '
     'Packages it all into a prompt the LLM understands.'),
    
    ('LLM Service',
     'Calls Nemotron via vLLM WebSocket or GPT-4 via OpenAI API. Passes the full context. '
     'Returns text + optional tool calls.'),
    
    ('Tool Execution Layer',
     'If LLM says "call verify_identity", the framework intercepts that, runs the Python function, '
     'gets the result, and feeds it back to the LLM.'),
    
    ('TTS Service',
     'Takes LLM output text (e.g., "Great! Your refill is ready.") and converts to audio. '
     'Gradium streams this as audio frames.'),
    
    ('Avatar Service (Optional)',
     'Takes TTS audio + transcript. Renders a synthetic face with lip-sync. '
     'Overlays on video stream (disabled on Twilio for low bandwidth).'),
]

for component, description in components:
    doc.add_heading(component, level=3)
    doc.add_paragraph(description)

doc.add_page_break()

# ============================================================================
# SECTION 8: TOOLS & VERIFICATION
# ============================================================================

doc.add_heading('8. The Tools System & Verification', level=1)

doc.add_heading('What are Tools?', level=2)
doc.add_paragraph(
    'Tools are Python async functions that the LLM can call when it needs to take action. '
    'The LLM doesn\'t run arbitrary code—it only calls pre-defined, fixed tools.'
)

doc.add_heading('The Four Tools in Bayview', level=2)

tools = [
    ('verify_identity(full_name, date_of_birth)',
     'Checks pharmacy records. Returns {verified: true/false}. '
     'MUST succeed before any other tools work. Hard security boundary.'),
    
    ('get_prescriptions(full_name)',
     'Returns list of caller\'s medications, refills left, pickup status. '
     'Only works if verified=true.'),
    
    ('refill_prescription(full_name, drug_name)',
     'Processes refill, checks refills remaining, updates mock_backend. '
     'Only works if verified=true.'),
    
    ('end_call()',
     'Gracefully terminates call. No parameters.'),
]

for tool_sig, tool_desc in tools:
    doc.add_heading(tool_sig, level=3)
    doc.add_paragraph(tool_desc)

doc.add_heading('How LLM Calls Tools', level=2)

tool_call_flow = [
    ('LLM receives user input: "Can I refill Lisinopril?"',
     'Plus system instruction + conversation history'),
    
    ('LLM outputs: "I need to verify identity first"',
     'And also emits: {function: "verify_identity", args: {full_name: "...", date_of_birth: "..."}}'),
    
    ('Pipecat framework intercepts the function call',
     'Does NOT let the LLM directly execute Python. Framework controls execution.'),
    
    ('Framework runs: await verify_identity(...)',
     'Inside bot-nemotron.py. Function queries mock_backend.py.'),
    
    ('Tool returns result: {verified: true}',
     'Framework feeds this back to LLM as a tool result message.'),
    
    ('LLM receives result and continues',
     'LLM can now call get_prescriptions since verify_identity succeeded.'),
]

for i, (step, detail) in enumerate(tool_call_flow, 1):
    doc.add_heading(f'Step {i}', level=3)
    doc.add_paragraph(step)
    doc.add_paragraph(detail, style='List Bullet')

doc.add_heading('The Verification Guardrail (Security by Code)', level=2)

doc.add_paragraph(
    'Here\'s what makes Bayview secure: verification is ENFORCED at the code level, not just the prompt level.'
)

doc.add_paragraph(
    '❌ Wrong approach: Tell the LLM "always verify first", hope it remembers.',
    style='List Bullet')

doc.add_paragraph(
    '✅ Right approach: Make get_prescriptions() and refill_prescription() check verified flag in code.',
    style='List Bullet')

doc.add_paragraph()

code_example = """# Inside get_prescriptions():
async def get_prescriptions(params, full_name):
    if not call_state["verified"]:
        return {
            "error": "Verify identity first",
            "note": "I must verify who you are before showing prescriptions"
        }
    # ... now safe to return prescriptions ...
    patient = find_patient_by_name(full_name)
    return {"prescriptions": patient["prescriptions"]}"""

code_para = doc.add_paragraph(code_example)
for run in code_para.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph(
    'This means NO prompt injection, trick question, or jailbreak can get the agent to leak prescription data. '
    'The defense is in the code.'
)

doc.add_page_break()

# ============================================================================
# SECTION 9: SELF-HEALING LOOP
# ============================================================================

doc.add_heading('9. Self-Healing Loop (Cekura)', level=1)

doc.add_heading('The Problem: Manual Testing', level=2)
doc.add_paragraph(
    'Traditionally, developers test an agent manually. But production always has edge cases. '
    'The Bayview solution: automate testing with Cekura, then automate fixes too.'
)

doc.add_heading('The Self-Healing Loop', level=2)

loop = [
    ('Cekura runs test scenarios against deployed agent',
     'Example: "Unverified caller requests prescriptions" (should be rejected)'),
    
    ('Test fails',
     'Agent incorrectly revealed prescription data before verification'),
    
    ('Cekura sends webhook to harness/webhook_server.py',
     'Includes the full transcript of the failure'),
    
    ('harness/self_heal.py receives webhook',
     'Deduplicates (same failure seen before?) and queues for repair'),
    
    ('GPT-5.5 proposes a fix',
     'self_heal.py asks: "This agent failed. Here\'s what went wrong. Suggest a code fix."'),
    
    ('Patch applied and redeployed',
     'The bot code is updated, pc cloud deploy runs'),
    
    ('Cekura re-runs the same test',
     'Does the agent now pass? Score improved?'),
    
    ('PR opens (if successful)',
     'Only merge if re-run score improves. No guess-and-check; every fix is measured.'),
]

for i, (step, detail) in enumerate(loop, 1):
    doc.add_heading(f'Step {i}: {step}', level=3)
    doc.add_paragraph(detail)

doc.add_heading('Key Harness Files', level=2)

harness_files = [
    ('webhook_server.py',
     'FastAPI server listening on /webhook. Receives Cekura failure notifications.'),
    
    ('self_heal.py',
     'Reads failures, calls GPT-5.5 with the transcript, applies patch to bot code.'),
    
    ('generate_dashboard.py',
     'Creates HTML dashboard showing KPIs, failure clusters, fix history.'),
    
    ('serve_dashboard.py',
     'Runs local web server (http://localhost:8765) with live Cekura data refresh.'),
]

for filename, description in harness_files:
    doc.add_heading(filename, level=3)
    doc.add_paragraph(description)

doc.add_page_break()

# ============================================================================
# SECTION 10: RUNNING LOCALLY
# ============================================================================

doc.add_heading('10. Running Locally', level=1)

doc.add_heading('Prerequisites', level=2)
doc.add_paragraph('Python 3.11+', style='List Bullet')
doc.add_paragraph('pip or uv (uv is recommended for this project)', style='List Bullet')
doc.add_paragraph('API keys in .env file (Gradium, OpenAI, NVIDIA endpoints)', style='List Bullet')
doc.add_paragraph('IMPORTANT: ENV=local in .env (required for local dev!)', style='List Bullet')

doc.add_heading('Step 1: Clone and Setup', level=2)

setup_commands = [
    'git clone https://github.com/quiet-node/yc-voice-agents-hackathon.git',
    'cd yc-voice-agents-hackathon/server',
    'cp .env.example .env  # Edit .env with your API keys',
]

for cmd in setup_commands:
    p = doc.add_paragraph(cmd)
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)

doc.add_heading('Step 2: Install Dependencies', level=2)

doc.add_paragraph('Using uv (recommended):')
cmd_para = doc.add_paragraph('uv sync')
cmd_para.runs[0].font.name = 'Courier New'

doc.add_paragraph()
doc.add_paragraph('Or using pip:')
cmd_para = doc.add_paragraph('pip install -r requirements.txt')
cmd_para.runs[0].font.name = 'Courier New'

doc.add_heading('Step 3: Start the Bot', level=2)

doc.add_paragraph('Using uv:')
cmd_para = doc.add_paragraph('uv run bot-nemotron.py')
cmd_para.runs[0].font.name = 'Courier New'

doc.add_paragraph()
doc.add_paragraph('Or using Python directly:')
cmd_para = doc.add_paragraph('python bot-nemotron.py')
cmd_para.runs[0].font.name = 'Courier New'

doc.add_paragraph()
doc.add_paragraph('You should see:')
doc.add_paragraph('INFO:     Started server process [1234]', style='List Bullet')
doc.add_paragraph('INFO:     Uvicorn running on http://0.0.0.0:7860', style='List Bullet')
doc.add_paragraph('INFO:     Application startup complete', style='List Bullet')

doc.add_heading('Step 4: Open the Pipecat Playground', level=2)

doc.add_paragraph('The Pipecat Playground is a built-in test UI at:')
cmd_para = doc.add_paragraph('http://localhost:7860')
cmd_para.runs[0].font.name = 'Courier New'

doc.add_paragraph()
doc.add_paragraph('Click "Connect" and the playground will stream audio/video to your bot. '
                 'You can now test the agent by speaking to it.')

doc.add_heading('Step 5: Test with Sample Patients', level=2)

test_patients = [
    ('Jane Doe', '1985-04-12', 'Lisinopril (2 refills, ready)'),
    ('John Smith', '1972-09-30', 'Metformin (5 refills, not ready)'),
    ('Maria Garcia', '1990-11-23', 'Levothyroxine + Albuterol'),
]

for name, dob, meds in test_patients:
    p = doc.add_paragraph(f'{name} / {dob}')
    p.style = 'List Bullet'
    doc.add_paragraph(f'Meds: {meds}', style='List Bullet 2')

doc.add_paragraph()
doc.add_paragraph('Try these conversations:')
doc.add_paragraph('Correct verification: "Hi, I\'m Jane Doe, born April 12th, 1985. Can I refill my Lisinopril?"', style='List Bullet')
doc.add_paragraph('Verify first: "Can I get my prescriptions?" (should be rejected)', style='List Bullet')
doc.add_paragraph('Spanish: "Hola, necesito una receta" or press "2" for Spanish', style='List Bullet')

doc.add_heading('Troubleshooting Local Dev', level=2)

doc.add_heading('"Address already in use" error', level=3)
doc.add_paragraph('Port 7860 is already taken. Kill the old process:')
cmd_para = doc.add_paragraph('lsof -i :7860  # Find process ID')
cmd_para.runs[0].font.name = 'Courier New'
doc.add_paragraph('kill -9 <PID>')
cmd_para = doc.paragraphs[-1].runs[0]
cmd_para.font.name = 'Courier New'

doc.add_heading('Imports fail (Krisp, VAD, etc.)', level=3)
doc.add_paragraph('Make sure ENV=local in .env. Without it, Pipecat tries to load '
                 'Pipecat Cloud-only features (Krisp noise filter) that aren\'t installed.')

doc.add_heading('Bot starts but doesn\'t respond', level=3)
doc.add_paragraph('Check:')
doc.add_paragraph('NEMOTRON_LLM_URL and NEMOTRON_ENABLE_THINKING are set correctly', style='List Bullet')
doc.add_paragraph('NVIDIA endpoints are up and accepting requests', style='List Bullet')
doc.add_paragraph('Fallback to bot-gpt.py if NVIDIA is unreliable (uses OpenAI GPT-4)', style='List Bullet')

doc.add_page_break()

# ============================================================================
# SECTION 11: DOCKER
# ============================================================================

doc.add_heading('11. Docker & Containerization', level=1)

doc.add_heading('Why Docker?', level=2)
doc.add_paragraph(
    'Docker packages the entire application (Python + dependencies + code) into a container. '
    'This ensures:'
)
doc.add_paragraph('Consistency: Runs the same way locally, in CI/CD, and in production', style='List Bullet')
doc.add_paragraph('Isolation: Doesn\'t interfere with other apps on the machine', style='List Bullet')
doc.add_paragraph('Reproducibility: No "works on my machine" surprises', style='List Bullet')
doc.add_paragraph('Deploying to Pipecat Cloud requires a Docker image', style='List Bullet')

doc.add_heading('The Dockerfile', level=2)

dockerfile_content = """FROM dailyco/pipecat-base:latest
    Base image: includes Pipecat, audio codecs, VAD models, etc.

ENV UV_COMPILE_BYTECODE=1
    Compile Python to bytecode for faster startup

RUN uv sync --locked
    Install Python dependencies from pyproject.toml + uv.lock
    (Pipecat Cloud caches this layer for speed)

COPY ./bot-nemotron.py bot.py
    Copy the main bot file (Pipecat Cloud runs bot.py)

COPY ./mock_backend.py mock_backend.py
COPY ./nemotron_llm.py nemotron_llm.py
COPY ./nvidia_stt.py nvidia_stt.py
...
    Copy all supporting Python modules"""

dockerfile_para = doc.add_paragraph(dockerfile_content)
for run in dockerfile_para.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_heading('Building the Docker Image Locally', level=2)

doc.add_paragraph('From the server/ directory:')

build_cmd = doc.add_paragraph('docker build -t bayview-pharmacy:latest .')
build_cmd.runs[0].font.name = 'Courier New'
build_cmd.runs[0].font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('This:')
doc.add_paragraph('Reads the Dockerfile', style='List Bullet')
doc.add_paragraph('Downloads the base image (pipecat-base)', style='List Bullet')
doc.add_paragraph('Installs Python dependencies', style='List Bullet')
doc.add_paragraph('Copies your code into the container', style='List Bullet')
doc.add_paragraph('Tags the image as bayview-pharmacy:latest', style='List Bullet')

doc.add_heading('Running a Docker Container Locally', level=2)

doc.add_paragraph('Basic run (for testing):')

docker_run = doc.add_paragraph(
    'docker run -p 7860:7860 --env-file .env bayview-pharmacy:latest'
)
docker_run.runs[0].font.name = 'Courier New'
docker_run.runs[0].font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('This:')
doc.add_paragraph('-p 7860:7860: Maps port 7860 (container) to 7860 (your machine)', style='List Bullet')
doc.add_paragraph('--env-file .env: Loads environment variables from .env file', style='List Bullet')
doc.add_paragraph('bayview-pharmacy:latest: Runs this image', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Interactive (with bash access for debugging):')

docker_interactive = doc.add_paragraph(
    'docker run -it -p 7860:7860 --env-file .env bayview-pharmacy:latest /bin/bash'
)
docker_interactive.runs[0].font.name = 'Courier New'
docker_interactive.runs[0].font.size = Pt(10)

doc.add_heading('Useful Docker Commands', level=2)

docker_commands = [
    ('docker build -t name:tag .', 'Build image from Dockerfile'),
    ('docker run -p 7860:7860 image:tag', 'Run container, map port'),
    ('docker ps', 'List running containers'),
    ('docker logs <CONTAINER_ID>', 'View container logs'),
    ('docker stop <CONTAINER_ID>', 'Stop a running container'),
    ('docker rm <CONTAINER_ID>', 'Remove a stopped container'),
    ('docker rmi image:tag', 'Delete an image'),
    ('docker exec -it <ID> bash', 'Open shell in running container'),
    ('docker push user/repo:tag', 'Push image to Docker Hub (for Cloud)'),
]

for cmd, desc in docker_commands:
    p = doc.add_paragraph(f'{cmd}')
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)
    doc.add_paragraph(desc, style='List Bullet 2')

doc.add_page_break()

# ============================================================================
# SECTION 12: DEPLOYING TO PRODUCTION
# ============================================================================

doc.add_heading('12. Deploying to Production', level=1)

doc.add_heading('Deployment Targets', level=2)

doc.add_paragraph('Pipecat Cloud: Primary (uses Daily.co for infrastructure, auto-scales)', style='List Bullet')
doc.add_paragraph('Custom VPS: Deploy Docker image to EC2, DigitalOcean, etc.', style='List Bullet')
doc.add_paragraph('Twilio: Phone gateway (uses WebSocket)', style='List Bullet')

doc.add_heading('Option 1: Deploy to Pipecat Cloud (Recommended)', level=2)

doc.add_paragraph('This is the easiest path for the hackathon.')

deployment_steps = [
    ('Install Pipecat CLI',
     'uv tool install pipecat-ai-cli'),
    
    ('Login',
     'pc cloud auth login (authenticates with your Pipecat Cloud account)'),
    
    ('Create secrets',
     'pc cloud secrets set config --file .env (uploads API keys securely)'),
    
    ('Deploy',
     'pc cloud deploy (builds Docker image, uploads to Cloud, deploys)'),
]

for i, (step, cmd) in enumerate(deployment_steps, 1):
    doc.add_heading(f'{i}. {step}', level=3)
    cmd_para = doc.add_paragraph(cmd)
    cmd_para.runs[0].font.name = 'Courier New'

doc.add_paragraph()
doc.add_paragraph('The bot will be live at a URL assigned by Pipecat Cloud (e.g., https://agent-xyz.pipecat.ai).')

doc.add_heading('Configuration: pcc-deploy.toml', level=2)

doc.add_paragraph(
    'Pipecat Cloud reads pcc-deploy.toml to know:'
)
doc.add_paragraph('Entry point: Which Python file to run (bot.py)', style='List Bullet')
doc.add_paragraph('Runtime: Python version, dependencies', style='List Bullet')
doc.add_paragraph('Secrets: Which environment variables are needed', style='List Bullet')

pcc_example = """[services]
bot = { run = "python bot.py" }

[secrets]
config = {}  # Reads from PIPECAT_CONFIG_VARS"""

pcc_para = doc.add_paragraph(pcc_example)
for run in pcc_para.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_heading('Option 2: Deploy to Custom Server (Advanced)', level=2)

doc.add_paragraph('If you want full control (or cheaper hosting):')

custom_deploy = [
    'Build Docker image locally: docker build -t mybot:latest .',
    'Push to registry: docker push registry.example.com/mybot:latest',
    'SSH into server: ssh user@your-server.com',
    'Pull image: docker pull registry.example.com/mybot:latest',
    'Run: docker run -d -p 7860:7860 --env-file .env mybot:latest',
    'Set up reverse proxy (nginx) pointing to localhost:7860',
]

for i, step in enumerate(custom_deploy, 1):
    p = doc.add_paragraph(step)
    p.style = 'List Number'

doc.add_heading('Option 3: Phone Gateway with Twilio', level=2)

doc.add_paragraph(
    'To accept PSTN phone calls:'
)

phone_steps = [
    'Get a Twilio account and phone number',
    'Create a TwiML Bin pointing to your agent URL',
    'Agent detects Twilio WebSocket connection (bot.py auto-switches transport)',
    'Calls routed through Twilio media gateway (8 kHz mu-law audio)',
    'Agent responds with voice (video disabled for phone)',
]

for i, step in enumerate(phone_steps, 1):
    p = doc.add_paragraph(step)
    p.style = 'List Number'

doc.add_page_break()

# ============================================================================
# SECTION 13: EXTENDING THE PROJECT
# ============================================================================

doc.add_heading('13. Extending the Project', level=1)

doc.add_heading('Common Customizations', level=2)

doc.add_heading('A. Adding a New Patient / Medication', level=3)
doc.add_paragraph('Edit server/mock_backend.py:')

add_patient_code = '''PATIENTS = {
    ("new patient name", "YYYY-MM-DD"): {
        "id": "p99",
        "prescriptions": [
            {
                "drug": "Your Medication Name",
                "refills_remaining": 5,
                "ready": True,
                "last_filled": "2026-05-20",
            },
        ],
    },
}'''

code_para = doc.add_paragraph(add_patient_code)
for run in code_para.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_heading('B. Changing the Personality', level=3)
doc.add_paragraph(
    'Edit the system_instruction in bot-nemotron.py (top of run_bot function). Change to:'
)
doc.add_paragraph('Be more formal/casual', style='List Bullet')
doc.add_paragraph('Speak in different accents or styles', style='List Bullet')
doc.add_paragraph('Add new behaviors (e.g., always remind about side effects)', style='List Bullet')

doc.add_heading('C. Adding a New Tool', level=3)
doc.add_paragraph('Example: "schedule_pickup" tool lets patients pick a time to collect meds:')

doc.add_paragraph(
    '1. Define the tool function in bot-nemotron.py:',
    style='List Number'
)
doc.add_paragraph(
    '''async def schedule_pickup(params, full_name, preferred_date):
    # Your code here''',
    style='List Bullet 2'
)

doc.add_paragraph('2. Register it:', style='List Number')
doc.add_paragraph('llm.register_direct_function(schedule_pickup)', style='List Bullet 2')

doc.add_paragraph('3. Add to ToolsSchema so LLM knows about it:', style='List Number')
doc.add_paragraph('{function_name: "schedule_pickup", parameters: {...}}', style='List Bullet 2')

doc.add_paragraph('4. Test locally, then deploy')

doc.add_heading('D. Connecting to a Real Database', level=3)
doc.add_paragraph(
    'Instead of mock_backend.py dict, query PostgreSQL, MySQL, or an API:'
)

db_steps = [
    'Install async database driver (e.g., asyncpg for PostgreSQL)',
    'Replace PATIENTS dict with async function that queries the database',
    'Update verify_identity(), get_prescriptions(), refill_prescription() to use database',
    'Handle errors: connection timeouts, 404s, permission errors',
]

for i, step in enumerate(db_steps, 1):
    p = doc.add_paragraph(step)
    p.style = 'List Number'

doc.add_heading('E. Adding Support for a New Language', level=3)
doc.add_paragraph('To add French, German, or another language:')

language_steps = [
    'Update language_router.py to detect the new language',
    'Update the system_instruction to mention the new language',
    'Test language selection in the UI (update menu: "1 for English, 2 for Spanish, 3 for French")',
    'Deploy and evaluate with Cekura',
]

for i, step in enumerate(language_steps, 1):
    p = doc.add_paragraph(step)
    p.style = 'List Number'

doc.add_heading('Advanced Feature Ideas', level=2)

doc.add_heading('Medication Interaction Checker', level=3)
doc.add_paragraph(
    'New tool: check_interactions(medication_list) → returns warnings if drugs interact dangerously.'
)

doc.add_heading('Refill History', level=3)
doc.add_paragraph(
    'New tool: get_refill_history(drug_name) → shows when the patient last refilled this med, and schedule.'
)

doc.add_heading('Dosage Reminders', level=3)
doc.add_paragraph(
    'After refill: "Take 1 tablet twice daily with food. Set a phone reminder?"'
)

doc.add_heading('Insurance Coverage Check', level=3)
doc.add_paragraph(
    'New tool: check_coverage(drug_name, insurance_id) → "This is covered. Your copay is $15."'
)

doc.add_heading('Appointment Scheduling', level=3)
doc.add_paragraph(
    'Let patients book pharmacist consultations: get_available_slots() → book_appointment()'
)

doc.add_heading('Testing Your Changes', level=2)

testing_workflow = [
    'Make code change',
    'Test locally (uv run bot-nemotron.py or python bot.py)',
    'Test in Pipecat Playground (http://localhost:7860)',
    'Run Cekura scenarios to check for regressions',
    'Monitor harness dashboard for failures',
    'Deploy to Pipecat Cloud: pc cloud deploy',
]

for i, step in enumerate(testing_workflow, 1):
    p = doc.add_paragraph(step)
    p.style = 'List Number'

doc.add_page_break()

# ============================================================================
# CONCLUSION
# ============================================================================

doc.add_heading('Conclusion', level=1)

doc.add_paragraph(
    'Bayview Pharmacy demonstrates a production-grade AI voice agent. By combining:'
)

doc.add_paragraph('Robust architecture (Pipecat pipeline with proper error handling)', style='List Bullet')
doc.add_paragraph('Frontend-backend separation with clear communication protocols', style='List Bullet')
doc.add_paragraph('Security by code (identity verification enforced in tool functions)', style='List Bullet')
doc.add_paragraph('Continuous improvement (Cekura self-healing loop)', style='List Bullet')
doc.add_paragraph('Accessible design (video, gestures, multilingual)', style='List Bullet')
doc.add_paragraph('Docker containerization for reproducible deployment', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph(
    'The result is an agent that can scale from local testing to production with confidence. '
    'Every failure is measured and fixed. Every deployment is tested. Security is non-negotiable.'
)

doc.add_heading('Quick Reference: Key Commands', level=2)

commands = [
    ('uv sync', 'Install dependencies'),
    ('uv run bot-nemotron.py', 'Start bot locally'),
    ('docker build -t bayview .',  'Build Docker image'),
    ('docker run -p 7860:7860 --env-file .env bayview', 'Run Docker container'),
    ('pc cloud auth login', 'Login to Pipecat Cloud'),
    ('pc cloud deploy', 'Deploy to Pipecat Cloud'),
]

for cmd, desc in commands:
    p = doc.add_paragraph(f'{cmd}')
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)
    doc.add_paragraph(desc, style='List Bullet 2')

doc.add_heading('Next Steps', level=2)

doc.add_paragraph('Read CLAUDE.md for hackathon conventions and gotchas', style='List Bullet')
doc.add_paragraph('Run locally and test with sample patients', style='List Bullet')
doc.add_paragraph('Explore the codebase: start with bot-nemotron.py', style='List Bullet')
doc.add_paragraph('Build Docker image and test: docker build && docker run', style='List Bullet')
doc.add_paragraph('Try a small customization (add a patient, change persona)', style='List Bullet')
doc.add_paragraph('Deploy to Pipecat Cloud and call the live agent', style='List Bullet')
doc.add_paragraph('Set up Cekura and harness for continuous improvement', style='List Bullet')

# ============================================================================
# SAVE DOCUMENT
# ============================================================================

output_path = 'Bayview_Pharmacy_Tutorial.docx'
doc.save(output_path)
print(f'✅ Enhanced tutorial saved to: {output_path}')
print()
print('📄 Document includes:')
print('   ✓ Project structure & file tree')
print('   ✓ Frontend architecture (demo_client, WebRTC, RTVI, MediaPipe)')
print('   ✓ Backend architecture (Pipecat, services, LLM, tools)')
print('   ✓ Frontend-backend communication flow')
print('   ✓ Docker build & run commands')
print('   ✓ Deployment to Pipecat Cloud')
print('   ✓ Complete beginner tutorials')
