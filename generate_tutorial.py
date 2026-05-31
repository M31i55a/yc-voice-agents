#!/usr/bin/env python3
"""
Generate a comprehensive tutorial document for the Bayview Pharmacy Voice Agent.
This script creates a Word document (.docx) explaining the project for beginners.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ============================================================================
# TITLE AND INTRODUCTION
# ============================================================================

title = doc.add_paragraph()
title_run = title.add_run('Bayview Pharmacy Voice Agent')
title_run.bold = True
title_run.font.size = Pt(24)
title_run.font.color.rgb = RGBColor(0, 102, 204)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('A Comprehensive Tutorial for Beginners')
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(14)
subtitle_run.italic = True
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

date_para = doc.add_paragraph('May 2026')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # Spacing

# ============================================================================
# TABLE OF CONTENTS
# ============================================================================

doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Introduction and Overview',
    '2. The Problem: Why This Project Exists',
    '3. The Solution: How Bayview Pharmacy Works',
    '4. Architecture Overview',
    '5. Deep Dive: The Code Structure',
    '6. Understanding the Pipeline',
    '7. The Tools System',
    '8. The Self-Healing Loop (Cekura)',
    '9. Running the Agent',
    '10. Extending the Project',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============================================================================
# SECTION 1: INTRODUCTION AND OVERVIEW
# ============================================================================

doc.add_heading('1. Introduction and Overview', level=1)

doc.add_heading('What is Bayview Pharmacy?', level=2)
doc.add_paragraph(
    'Bayview Pharmacy is a voice-controlled AI agent that helps patients refill their prescriptions '
    'and check medication status. Instead of calling a pharmacy and waiting on hold, patients can have '
    'a conversation with an AI assistant that understands their needs, verifies their identity, and '
    'processes refill requests in real-time.'
)

doc.add_paragraph(
    'Think of it like this: imagine calling your pharmacy, but instead of waiting 5-10 minutes on hold '
    'to talk to a busy pharmacist, you immediately talk to an intelligent AI assistant that:'
)

doc.add_paragraph('Listens to your request using speech-to-text (STT)', style='List Bullet')
doc.add_paragraph('Understands what you want using an advanced language model (LLM)', style='List Bullet')
doc.add_paragraph('Checks your prescription records securely', style='List Bullet')
doc.add_paragraph('Processes your refill or provides information', style='List Bullet')
doc.add_paragraph('Speaks back to you with natural-sounding speech synthesis (TTS)', style='List Bullet')

doc.add_heading('Key Features', level=2)
doc.add_paragraph('🎥 Video Support: Calls include a visible AI avatar so users can see the agent they\'re talking to.',
                 style='List Bullet')
doc.add_paragraph('🗣️ Multilingual: English and Spanish support with automatic language detection.',
                 style='List Bullet')
doc.add_paragraph('🔐 Secure: Verifies identity before revealing any prescription information.',
                 style='List Bullet')
doc.add_paragraph('🚀 Always Improving: Uses Cekura to evaluate performance and automatically fix issues.',
                 style='List Bullet')
doc.add_paragraph('📱 Works Everywhere: Phone calls (Twilio), video web calls (WebRTC), and local testing.',
                 style='List Bullet')

doc.add_page_break()

# ============================================================================
# SECTION 2: THE PROBLEM
# ============================================================================

doc.add_heading('2. The Problem: Why This Project Exists', level=1)

doc.add_heading('The Challenge of Prescription Management', level=2)
doc.add_paragraph(
    'Millions of people struggle with prescription refills. Here\'s why it\'s a real problem:'
)

doc.add_heading('Pain Points for Patients:', level=3)
p = doc.add_paragraph('Long wait times on hold listening to hold music')
p.style = 'List Bullet'

p = doc.add_paragraph('Repeating sensitive information (name, date of birth) multiple times')
p.style = 'List Bullet'

p = doc.add_paragraph('Difficulty understanding medication instructions over phone calls')
p.style = 'List Bullet'

p = doc.add_paragraph('Accessibility issues for elderly patients, people with hearing loss, or speech difficulties')
p.style = 'List Bullet'

p = doc.add_paragraph('No way to check prescription status 24/7 outside business hours')
p.style = 'List Bullet'

doc.add_heading('Who Suffers Most?', level=3)
doc.add_paragraph('Elderly patients: Difficulty hearing, seeing, or speaking clearly during calls')
p = doc.add_paragraph('Stroke survivors and people with speech disabilities: Hard to communicate over phone')
p.style = 'List Bullet'

p = doc.add_paragraph('Non-English speakers: Language barriers with pharmacy staff')
p.style = 'List Bullet'

p = doc.add_paragraph('Working professionals: Can\'t call during business hours')
p.style = 'List Bullet'

p = doc.add_paragraph('People with hearing loss: Phone calls are difficult or impossible')
p.style = 'List Bullet'

doc.add_heading('The Business Problem:', level=3)
doc.add_paragraph(
    'Pharmacies are understaffed. Prescription refill calls take time that could be spent on clinical '
    'work. This creates a bottleneck where pharmacies can\'t serve all their patients efficiently. '
    'An AI agent can handle routine refill requests 24/7, freeing up staff for complex cases.'
)

doc.add_page_break()

# ============================================================================
# SECTION 3: THE SOLUTION
# ============================================================================

doc.add_heading('3. The Solution: How Bayview Pharmacy Works', level=1)

doc.add_heading('The User Experience', level=2)
doc.add_paragraph(
    'Here\'s what happens when a patient calls or video chats with Bayview:'
)

# Create a numbered list of steps
steps = [
    ('Patient initiates a call or starts a video session',
     'They call a phone number or click a link to a web app'),
    
    ('Agent greets them with language selection',
     'Agent asks "Press 1 for English, 2 for Spanish" and can detect preference from speech'),
    
    ('Agent requests identity verification',
     'Asks for full name and date of birth (this is MANDATORY before revealing any prescription data)'),
    
    ('Agent verifies information securely',
     'Checks the name and date of birth against the pharmacy records'),
    
    ('Patient can now request refills or check status',
     'Examples: "Can I refill my Lisinopril?", "Is my medication ready for pickup?"'),
    
    ('Agent processes the refill or provides information',
     'Confirms details and updates the prescription records'),
    
    ('Patient gets confirmation',
     'Hears when the prescription will be ready for pickup'),
]

for i, (step, detail) in enumerate(steps, 1):
    p = doc.add_paragraph(f'Step {i}: {step}')
    p.style = 'List Number'
    p_detail = doc.add_paragraph(detail, style='List Bullet 2')

doc.add_heading('Why This Design?', level=2)
doc.add_paragraph(
    'Identity verification is enforced in CODE, not just suggested in the prompt. This means:'
)
doc.add_paragraph('The agent CANNOT reveal prescription info before verifying identity', style='List Bullet')
doc.add_paragraph('If a caller refuses to give their DOB, the agent won\'t proceed', style='List Bullet')
doc.add_paragraph('Security is built into the system, not relying on the AI being "nice"', style='List Bullet')

doc.add_page_break()

# ============================================================================
# SECTION 4: ARCHITECTURE
# ============================================================================

doc.add_heading('4. Architecture Overview', level=1)

doc.add_paragraph(
    'Bayview is built on a modular pipeline architecture. Think of it like a factory assembly line '
    'where each station does one specific job:'
)

# ASCII Art-style description
doc.add_paragraph('Caller Audio → STT → LLM (+ Tools) → TTS → Caller Audio')
p = doc.paragraphs[-1]
p.runs[0].font.name = 'Courier New'
p.runs[0].font.size = Pt(10)

doc.add_heading('The Pipeline Stages:', level=2)

stages = [
    ('STT (Speech-to-Text)',
     'Gradium or NVIDIA Parakeet',
     'Converts the caller\'s voice into text that the LLM can understand'),
    
    ('Language Router',
     'Custom Python processor',
     'Detects if the caller wants English or Spanish and injects that context'),
    
    ('LLM (Large Language Model)',
     'Nemotron 3 Super (NVIDIA) or GPT-4',
     'The "brain" that understands requests, decides which tools to use, and generates responses'),
    
    ('Tools',
     'Python async functions',
     'Specific actions the LLM can call: verify_identity, get_prescriptions, refill_prescription, end_call'),
    
    ('TTS (Text-to-Speech)',
     'Gradium TTS',
     'Converts the LLM\'s response back into natural-sounding speech the caller hears'),
    
    ('Optional: Video Avatar',
     'AI-generated video overlay',
     'Shows a realistic avatar face with lip sync on video calls'),
]

for stage, service, description in stages:
    doc.add_heading(stage, level=3)
    doc.add_paragraph(f'Service: {service}')
    doc.add_paragraph(f'Purpose: {description}')

doc.add_page_break()

# ============================================================================
# SECTION 5: DEEP DIVE - CODE STRUCTURE
# ============================================================================

doc.add_heading('5. Deep Dive: The Code Structure', level=1)

doc.add_paragraph(
    'The code is organized in the `server/` directory. Let\'s look at each important file:'
)

doc.add_heading('bot-nemotron.py (The Main Agent)', level=2)
doc.add_paragraph(
    'This is the PRIMARY entry point. It\'s the brain of the entire operation. About 500-600 lines of Python that:'
)
doc.add_paragraph('Imports all the services (STT, LLM, TTS)', style='List Bullet')
doc.add_paragraph('Defines the LLM system instruction (persona, rules, behaviors)', style='List Bullet')
doc.add_paragraph('Implements the four main tools as Python async functions', style='List Bullet')
doc.add_paragraph('Manages the call state (verification status, refill history, etc.)', style='List Bullet')
doc.add_paragraph('Orchestrates the pipeline with Pipecat', style='List Bullet')

doc.add_paragraph()
doc.add_heading('Key Components Inside bot-nemotron.py:', level=3)

doc.add_paragraph(
    '🔐 System Instruction:\n'
    'A long prompt that tells the LLM how to behave. It says things like:'
)
doc.add_paragraph('"You are a helpful pharmacy assistant"', style='List Bullet')
doc.add_paragraph('"ALWAYS verify identity first"', style='List Bullet')
doc.add_paragraph('"Be empathetic and clear"', style='List Bullet')
doc.add_paragraph('"Support both English and Spanish"', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('📝 Tool Functions:\n'
                 'Four async functions that the LLM can "call" when it needs to do something:')

doc.add_paragraph()
tool_definitions = [
    ('verify_identity(full_name, date_of_birth)',
     'Checks if the caller\'s name and DOB match pharmacy records. MUST succeed before any other tools work.'),
    
    ('get_prescriptions(full_name)',
     'Returns a list of the caller\'s medications, refills remaining, and pickup status.'),
    
    ('refill_prescription(full_name, drug_name)',
     'Processes a refill request. Checks if refills are available and marks the prescription ready.'),
    
    ('end_call()',
     'Gracefully ends the conversation.'),
]

for tool_sig, tool_desc in tool_definitions:
    doc.add_heading(tool_sig, level=4)
    doc.add_paragraph(tool_desc)

doc.add_heading('mock_backend.py (The "Database")', level=2)
doc.add_paragraph(
    'A simple Python dictionary that stores patient data. In a real system, this would query a '
    'database, but for the hackathon it\'s in-memory:'
)
doc.add_paragraph('Patient records: name, date of birth', style='List Bullet')
doc.add_paragraph('For each patient: list of prescriptions', style='List Bullet')
doc.add_paragraph('For each prescription: drug name, refills remaining, ready status', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph(
    'This is the FILE TO EDIT if you want to add new patients or change prescription data. '
    'To connect to a real database (MySQL, PostgreSQL, etc.), you\'d replace this dictionary with API calls.'
)

doc.add_heading('language_router.py (Multilingual Support)', level=2)
doc.add_paragraph(
    'A custom processor that detects whether the caller wants English or Spanish. It looks for:'
)
doc.add_paragraph('Explicit keywords: "English", "Spanish", "Español", "Hola"', style='List Bullet')
doc.add_paragraph('Menu selections: caller presses 1 for English, 2 for Spanish', style='List Bullet')
doc.add_paragraph('And injects language preference into the LLM context', style='List Bullet')

doc.add_heading('nemotron_llm.py (LLM Service)', level=2)
doc.add_paragraph(
    'Wraps the NVIDIA Nemotron language model in a way Pipecat understands. '
    'If NVIDIA services are unavailable, the system falls back to OpenAI GPT-4 (bot-gpt.py uses this).'
)

doc.add_heading('nvidia_stt.py (Custom Speech-to-Text)', level=2)
doc.add_paragraph(
    'A custom WebSocket client for NVIDIA\'s Parakeet STT. This is optional; Gradium STT is the default. '
    'You\'d use this when NVIDIA STT is available and you want the open-weights advantage.'
)

doc.add_heading('video_avatar.py (Optional AI Avatar)', level=2)
doc.add_paragraph(
    'Integrates Pipecat\'s video avatar layer. If enabled, the caller sees a realistic AI face '
    'with lip-synced speech. On phone calls, this is disabled automatically.'
)

doc.add_page_break()

# ============================================================================
# SECTION 6: THE PIPELINE IN DETAIL
# ============================================================================

doc.add_heading('6. Understanding the Pipeline', level=1)

doc.add_heading('What is a Pipeline?', level=2)
doc.add_paragraph(
    'A pipeline is a series of processors that handle frames (units of data) sequentially. '
    'In Bayview:'
)

doc.add_paragraph('Audio frames come in from the caller', style='List Bullet')
doc.add_paragraph('They flow through STT (audio → text)', style='List Bullet')
doc.add_paragraph('Then through the LLM (text → action)', style='List Bullet')
doc.add_paragraph('Then back through TTS (text → audio)', style='List Bullet')
doc.add_paragraph('And audio frames go back out to the caller', style='List Bullet')

doc.add_heading('The Call State Machine', level=2)
doc.add_paragraph('Every call has state that persists across multiple messages:')

doc.add_paragraph('verified: Boolean flag. True only after verify_identity() succeeds.', style='List Bullet')
doc.add_paragraph('verified_name: The caller\'s name after verification.', style='List Bullet')
doc.add_paragraph('failed_attempts: Counter for failed identity checks.', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph(
    'This state is ISOLATED per call. If two callers are on the line simultaneously, '
    'each gets their own copy of the state. This is critical for security and privacy.'
)

doc.add_heading('Error Handling & Robustness', level=2)
doc.add_paragraph(
    'The agent is built to handle edge cases:'
)

doc.add_paragraph('Mispronounced names: Normalized and matched tolerantly ("Jane Do" matches "Jane Doe")',
                 style='List Bullet')
doc.add_paragraph('Spoken dates: Parsed robustly (e.g., "April twelfth, nineteen eighty-five" → "1985-04-12")',
                 style='List Bullet')
doc.add_paragraph('Failed verification: After 2 attempts, hands off to a human pharmacist',
                 style='List Bullet')
doc.add_paragraph('Missing drugs: Gracefully reports if a medication isn\'t found',
                 style='List Bullet')

doc.add_page_break()

# ============================================================================
# SECTION 7: THE TOOLS SYSTEM
# ============================================================================

doc.add_heading('7. The Tools System', level=1)

doc.add_heading('How Does the LLM Call Tools?', level=2)
doc.add_paragraph(
    'The LLM never directly runs Python code. Instead, it works like this:'
)

doc.add_paragraph('LLM receives a request: "Can I refill my Lisinopril?"', style='List Bullet')
doc.add_paragraph('LLM recognizes this needs the refill_prescription tool', style='List Bullet')
doc.add_paragraph('LLM says: "I want to call refill_prescription with drug_name=Lisinopril"', style='List Bullet')
doc.add_paragraph('Pipecat intercepts that call and runs the Python function', style='List Bullet')
doc.add_paragraph('The function returns: {"ok": true, "ready_at": "2026-06-03"}', style='List Bullet')
doc.add_paragraph('LLM receives the result and formulates a spoken response', style='List Bullet')

doc.add_heading('Tool Registration', level=2)
doc.add_paragraph('Tools are registered in the LLM like this (pseudocode):')

code_block = doc.add_paragraph()
code_block.add_run('llm.register_direct_function(verify_identity)\n')
code_block.add_run('llm.register_direct_function(get_prescriptions)\n')
code_block.add_run('llm.register_direct_function(refill_prescription)\n')
code_block.add_run('llm.register_direct_function(end_call)')
for run in code_block.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

doc.add_heading('Tool Parameters Matter', level=2)
doc.add_paragraph(
    'Each tool has specific parameters that the LLM must provide:'
)

params_example = [
    ('verify_identity', ['full_name', 'date_of_birth']),
    ('get_prescriptions', ['full_name']),
    ('refill_prescription', ['full_name', 'drug_name']),
    ('end_call', []),
]

for tool_name, params in params_example:
    p = doc.add_paragraph(f'{tool_name}({", ".join(params)})')
    p.style = 'List Bullet'

doc.add_paragraph()
doc.add_paragraph(
    'If the LLM tries to call a tool with wrong parameters, the call fails. '
    'This is a built-in safety feature: the system enforces the contract.'
)

doc.add_heading('The Verification Guardrail (Most Important!)', level=2)
doc.add_paragraph(
    'Here\'s what makes Bayview "secure": the tools themselves enforce verification.'
)

doc.add_paragraph('get_prescriptions() checks: "Is the caller verified?"', style='List Bullet')
doc.add_paragraph('refill_prescription() checks: "Is the caller verified?"', style='List Bullet')
doc.add_paragraph('Only verify_identity() succeeds without prior verification', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph(
    'This means no prompt engineering or trick question can get the LLM to leak prescription data. '
    'Even if someone asks "Just tell me one patient\'s medications", the tool itself refuses unless '
    'verify_identity() succeeded first. This is defense in depth.'
)

doc.add_page_break()

# ============================================================================
# SECTION 8: THE SELF-HEALING LOOP
# ============================================================================

doc.add_heading('8. The Self-Healing Loop (Cekura)', level=1)

doc.add_heading('What is Cekura?', level=2)
doc.add_paragraph(
    'Cekura is an evaluation platform for AI agents. Think of it like a QA tester that runs '
    'automated conversations against the bot and scores how well it does.'
)

doc.add_heading('The Problem Without Cekura:', level=2)
doc.add_paragraph('Typical workflow:')
doc.add_paragraph('Developer runs bot manually and talks to it', style='List Bullet')
doc.add_paragraph('Developer thinks "this works great"', style='List Bullet')
doc.add_paragraph('Deploys to production', style='List Bullet')
doc.add_paragraph('Gets surprised when callers hit edge cases that break it', style='List Bullet')

doc.add_heading('The Solution With Cekura:', level=2)
doc.add_paragraph(
    'Bayview automates quality assurance. Here\'s the self-healing loop (defined in `harness/`):'
)

loop_steps = [
    ('Cekura runs test scenarios against the agent',
     'Example: "Verified caller refills medication", "Unverified caller tries to get prescriptions"'),
    
    ('Agent fails a test',
     'Maybe it reveals prescription data before verifying identity'),
    
    ('Webhook fires',
     'Cekura sends a failure notification to webhook_server.py'),
    
    ('GPT-5.5 proposes a patch',
     'The harness asks: "The agent failed this test. Here\'s the transcript. How should we fix the bot?"'),
    
    ('Patch is applied and deployed',
     'The bot code is updated and redeployed to Pipecat Cloud'),
    
    ('Cekura re-runs the same test',
     'Does the fix work? Did the score improve?'),
    
    ('PR opens (if successful)',
     'Only if the re-run score improves do we open a pull request to merge the fix'),
]

for i, (step, detail) in enumerate(loop_steps, 1):
    p = doc.add_paragraph(f'{i}. {step}')
    p.style = 'List Number'
    p_detail = doc.add_paragraph(detail, style='List Bullet 2')

doc.add_heading('Key Files in the Harness:', level=2)

harness_files = [
    ('self_heal.py',
     'Asks GPT-5.5 for a targeted patch when a test fails'),
    
    ('webhook_server.py',
     'Listens for failure notifications from Cekura'),
    
    ('generate_dashboard.py',
     'Creates an interactive dashboard showing KPIs, failures, and fix status'),
    
    ('serve_dashboard.py',
     'Runs a web server that auto-refreshes with latest Cekura data'),
]

for filename, purpose in harness_files:
    doc.add_heading(filename, level=3)
    doc.add_paragraph(purpose)

doc.add_heading('Why This Matters:', level=2)
doc.add_paragraph(
    'Traditional agents are like released software: developers hope they\'re good enough, but bugs appear in production. '
    'The Bayview self-healing loop means:'
)
doc.add_paragraph('Every failure becomes a regression test', style='List Bullet')
doc.add_paragraph('Fixes are measured, not guessed', style='List Bullet')
doc.add_paragraph('Only improvements get merged', style='List Bullet')
doc.add_paragraph('The agent gets continuously smarter', style='List Bullet')

doc.add_page_break()

# ============================================================================
# SECTION 9: RUNNING THE AGENT
# ============================================================================

doc.add_heading('9. Running the Agent', level=1)

doc.add_heading('Prerequisites', level=2)
doc.add_paragraph('Python 3.11 or newer')
doc.add_paragraph('A `.env` file with API keys', style='List Bullet')

doc.add_heading('Environment Variables You Need (.env)', level=2)
doc.add_paragraph('Copy `server/.env.example` to `server/.env` and fill in:')

env_table_data = [
    ['Variable', 'For', 'Example'],
    ['GRADIUM_API_KEY', 'TTS', 'Your Gradium API key'],
    ['GRADIUM_VOICE_ID', 'TTS', 'en_US_mary'],
    ['OPENAI_API_KEY', 'Fallback LLM', 'sk-...'],
    ['NVIDIA_ASR_URL', 'Nemotron STT', 'ws://nvidia-asr-service'],
    ['NEMOTRON_LLM_URL', 'Nemotron LLM', 'http://nemotron-service'],
    ['NEMOTRON_LLM_MODEL', 'Nemotron version', 'nemotron-3-super-120b'],
    ['ENV', 'For local dev', 'local (REQUIRED!)'],
]

# Add as paragraph list for simpler document
for row in env_table_data[1:]:
    p = doc.add_paragraph(f'{row[0]}: {row[1]}')
    p.style = 'List Bullet'

doc.add_heading('Quickstart: Local Development', level=2)
doc.add_paragraph('From the `server/` directory:')

quickstart_commands = [
    'uv sync  # Install dependencies (first time only)',
    'uv run bot-nemotron.py  # Start the bot',
]

for cmd in quickstart_commands:
    p = doc.add_paragraph(cmd)
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph(
    'Open http://localhost:7860 (the Pipecat Playground) and click **Connect**. '
    'You can now talk to the bot for testing.'
)

doc.add_heading('Test Accounts (Mock Data)', level=2)
doc.add_paragraph('Use these to test:')

test_accounts = [
    ('Jane Doe', '1985-04-12', 'Lisinopril 10mg (2 refills, ready)'),
    ('John Smith', '1972-09-30', 'Metformin 500mg (5 refills, not ready)'),
    ('Maria Garcia', '1990-11-23', 'Levothyroxine + Albuterol'),
]

for name, dob, meds in test_accounts:
    p = doc.add_paragraph(f'{name} / {dob}')
    p.style = 'List Bullet'
    doc.add_paragraph(f'Medications: {meds}', style='List Bullet 2')

doc.add_heading('Deploying to Pipecat Cloud', level=2)
doc.add_paragraph('Once tested locally:')

deploy_steps = [
    'Install: uv tool install pipecat-ai-cli',
    'Login: pc cloud auth login',
    'Upload secrets: pc cloud secrets set config --file .env',
    'Deploy: pc cloud deploy',
]

for i, step in enumerate(deploy_steps, 1):
    p = doc.add_paragraph(step)
    p.style = 'List Number'

doc.add_page_break()

# ============================================================================
# SECTION 10: EXTENDING THE PROJECT
# ============================================================================

doc.add_heading('10. Extending the Project', level=1)

doc.add_heading('Common Customizations', level=2)

doc.add_heading('A. Adding a New Patient', level=3)
doc.add_paragraph('Edit `mock_backend.py`:')

sample_code = doc.add_paragraph()
sample_code.add_run('PATIENTS = {\n')
sample_code.add_run('    ("new patient", "1980-05-15"): {\n')
sample_code.add_run('        "id": "p6",\n')
sample_code.add_run('        "prescriptions": [\n')
sample_code.add_run('            {"drug": "Your Med", "refills_remaining": 3, "ready": True, "last_filled": "2026-05-20"},\n')
sample_code.add_run('        ],\n')
sample_code.add_run('    },\n')
sample_code.add_run('}')
for run in sample_code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_heading('B. Changing the System Instruction', level=3)
doc.add_paragraph(
    'Edit the `system_instruction` variable in `bot-nemotron.py`. This is the "personality" '
    'and ruleset for the agent. Change it to:'
)
doc.add_paragraph('Sound more/less formal', style='List Bullet')
doc.add_paragraph('Add new personality traits', style='List Bullet')
doc.add_paragraph('Change security policies (e.g., require more information before refills)', style='List Bullet')
doc.add_paragraph('Support new languages (update language_router.py too)', style='List Bullet')

doc.add_heading('C. Adding a New Tool', level=3)
doc.add_paragraph('To let the LLM perform a new action (e.g., schedule a refill pickup):')

new_tool_steps = [
    'Write a new async function in bot-nemotron.py:',
    'Register it with llm.register_direct_function()',
    'Add it to the ToolsSchema (the LLM\'s knowledge of available tools)',
    'Test it locally using the Playground',
    'Deploy and evaluate with Cekura',
]

for i, step in enumerate(new_tool_steps, 1):
    p = doc.add_paragraph(step)
    p.style = 'List Number'

doc.add_heading('D. Connecting to a Real Database', level=3)
doc.add_paragraph(
    'Instead of using the mock `PATIENTS` dict, query a real database:'
)

db_steps = [
    'Install a database library (e.g., asyncpg for PostgreSQL)',
    'Replace the mock_backend.py dict with async database queries',
    'Update verify_identity(), get_prescriptions(), and refill_prescription() to use real queries',
    'Handle database errors gracefully (network failures, timeouts, etc.)',
]

for i, step in enumerate(db_steps, 1):
    p = doc.add_paragraph(step)
    p.style = 'List Number'

doc.add_heading('E. Changing the TTS Voice', level=3)
doc.add_paragraph(
    'Edit the `GRADIUM_VOICE_ID` in `.env`. Different voice IDs produce different voices '
    '(male, female, accent, etc.). Check Gradium\'s documentation for available voices.'
)

doc.add_heading('Advanced Extensions:', level=2)

doc.add_heading('New Feature: Medication Interactions Checker', level=3)
doc.add_paragraph(
    'Add a new tool that checks if medications interact dangerously:'
)

doc.add_paragraph('Tool name: check_interactions', style='List Bullet')
doc.add_paragraph('Parameters: list of drug names', style='List Bullet')
doc.add_paragraph('Returns: warnings if drugs don\'t mix well', style='List Bullet')
doc.add_paragraph('Call it before refill approval', style='List Bullet')

doc.add_heading('New Feature: Appointment Scheduling', level=3)
doc.add_paragraph(
    'Add tools to let patients book pharmacist consultations:'
)

doc.add_paragraph('Tool: get_available_slots() - returns open appointment times', style='List Bullet')
doc.add_paragraph('Tool: book_appointment() - schedules a consultation', style='List Bullet')

doc.add_heading('New Feature: Multi-Channel Support', level=3)
doc.add_paragraph(
    'Extend beyond phone and WebRTC:'
)

doc.add_paragraph('SMS refill requests (via Twilio SMS)', style='List Bullet')
doc.add_paragraph('Email confirmations', style='List Bullet')
doc.add_paragraph('WhatsApp integration', style='List Bullet')

doc.add_heading('Testing Your Changes', level=2)
doc.add_paragraph('After any code change:')

testing_steps = [
    'Run bot-nemotron.py locally',
    'Test through the Pipecat Playground',
    'Use Cekura to run automated scenarios',
    'Monitor the harness dashboard for failures',
    'Deploy to Pipecat Cloud only when tests pass',
]

for i, step in enumerate(testing_steps, 1):
    p = doc.add_paragraph(step)
    p.style = 'List Number'

doc.add_page_break()

# ============================================================================
# CONCLUSION
# ============================================================================

doc.add_heading('Conclusion', level=1)

doc.add_paragraph(
    'Bayview Pharmacy demonstrates how AI agents can solve real healthcare problems. '
    'By combining:'
)

doc.add_paragraph('Robust voice interfaces (STT + LLM + TTS)', style='List Bullet')
doc.add_paragraph('Enforced security guarantees (identity verification in code)', style='List Bullet')
doc.add_paragraph('Continuous improvement loops (Cekura self-healing)', style='List Bullet')
doc.add_paragraph('Accessibility-first design (video, gestures, multilingual)', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph(
    'The project shows that production-grade AI agents aren\'t guesswork—they\'re systems where '
    'every failure drives measurable improvement, security is non-negotiable, and the whole pipeline '
    'is transparent and debuggable.'
)

doc.add_heading('Next Steps:', level=2)
doc.add_paragraph('Read CLAUDE.md (the hackathon runbook) for additional context', style='List Bullet')
doc.add_paragraph('Clone the repo and run it locally', style='List Bullet')
doc.add_paragraph('Experiment with mock_backend.py to add your own patients', style='List Bullet')
doc.add_paragraph('Call +1 (628) 300-0587 to experience the live agent', style='List Bullet')
doc.add_paragraph('Review the self-heal harness to understand continuous improvement', style='List Bullet')

doc.add_heading('Key Takeaways:', level=2)
doc.add_paragraph('AI agents need structure, not just intelligence', style='List Bullet')
doc.add_paragraph('Security must be code-level, not prompt-level', style='List Bullet')
doc.add_paragraph('Evaluation drives improvement (measure, don\'t guess)', style='List Bullet')
doc.add_paragraph('Accessibility benefits everyone (video, multilingual, clear voice)', style='List Bullet')
doc.add_paragraph('Production agents need robust error handling and fallbacks', style='List Bullet')

# ============================================================================
# SAVE DOCUMENT
# ============================================================================

output_path = 'Bayview_Pharmacy_Tutorial.docx'
doc.save(output_path)
print(f'✅ Tutorial saved to: {output_path}')
print(f'📄 Document created successfully with comprehensive explanations for beginners!')
